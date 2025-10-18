import glob
import os
import shutil
import tkinter as tk
import unicodedata
from io import BytesIO
from pathlib import Path
from tkinter import filedialog
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

import pandas as pd
from PIL import Image
from docx import Document
from docx.oxml.ns import qn


def get_files(dir_, ending):
    pdfs = []
    dir_ = os.path.abspath(dir_)
    for filename in os.listdir(dir_):
        if filename.endswith(ending):
            full_path = os.path.join(dir_, filename)
            pdfs.append(full_path)
        else:
            continue
    return pdfs


def get_file():
    is_valid = False
    path = ""
    while not is_valid:
        path = str(input("Bitte gib den Pfad zu einer Datei (.xlsx, .csv, .docx) ein:"))
        path = path.replace("\"", "")
        path = path.replace("'", "")
        if path.endswith(".xlsx") or path.endswith(".csv") or path.endswith(".docx"):
            is_valid = True
    path = os.path.abspath(path)
    return path


def get_path():
    is_valid = False
    path = ""
    while not is_valid:
        path = str(input("Bitte gib den Pfad zu einem Order ein:"))
        if os.path.exists(path):
            is_valid = True
    path = os.path.abspath(path)
    return path


def read_file(path, images):
    df = pd.DataFrame()
    imgs = {}
    try:
        if path.endswith(".xlsx"):
            df = read_excel(path)
        elif path.endswith(".csv"):
            df = pd.read_csv(path, dtype=str, keep_default_na=False)
        elif path.endswith(".docx"):
            if images:
                df, imgs = read_docx_with_images(path)
            else:
                df = read_docx(path)
    except ValueError:
        print("Datei ist ungültig")
        return df, imgs
    except FileNotFoundError:
        print(f"Datei {path} existiert nicht")
        return df, imgs
    except Exception as err:
        print("Etwas anderes ist schief gelaufen", err)
        return df, imgs
    return df, imgs


def get_several_files_ui():
    root = tk.Tk()
    root.withdraw()
    file_paths = filedialog.askopenfilenames(
        title="Wähle mehrere Word- oder Exceltabellen aus",
        filetypes=[("Word- und Excel-Dateien", "*.docx *.xlsx *.xls")]
    )
    return list(file_paths)


def get_file_ui():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="Wähle eine Datei aus",
        filetypes=[("Word-,CSV- oder Excel-Datei", "*.docx *.xlsx *.xls *.csv")]
    )
    return file_path


def read_files(files, images):
    res = []
    imgs = []
    for file in files:
        tmp, img_dict = read_file(file, images)
        if not tmp.empty:
            res.append(tmp)
            imgs.append(img_dict)
        else:
            print("Konnte Datei ", file, "nicht lesen")
    return res, imgs


def iter_unique_cells(cells):
    prior_cell = None
    for c in cells:
        if c == prior_cell:
            continue
        yield c
        prior_cell = c


def read_docx(path):
    document = Document(path)
    data = []
    for table in document.tables:
        for row in table.rows:
            data.append([cell.text.strip() for cell in list(iter_unique_cells(row.cells))])

    # Convert the data to a DataFrame
    df = pd.DataFrame(data=data, dtype=str)

    # Optional: If the first row is the header
    df.columns = df.iloc[0]
    df = df[1:].reset_index(drop=True)
    return df


def read_excel(file):
    # Read first 20 rows without headers to detect the header row
    preview = pd.read_excel(file, header=None, nrows=20)

    # Find first row index where there's at least one non-empty cell (likely header)
    header_row_idx = preview.apply(lambda row: row.notna().any(), axis=1).idxmax()

    # Read again using that row as header, skip everything before
    df = pd.read_excel(file, dtype=str, keep_default_na=False, header=header_row_idx)

    # Optional: drop fully empty rows below the table if any
    df = df.dropna(how='all')
    return df


def _extract_cell(cell):
    pictures = []
    paragraph_texts = []

    for para in cell.paragraphs:
        fragments = []
        for run in para.runs:
            run_elem = run._element
            for node in run_elem.iter():
                if node.tag.endswith('}t') and node.text:
                    fragments.append(node.text)
                elif node.tag.endswith('}br'):
                    fragments.append('\n')  # soft line break

        paragraph_text = ''.join(fragments)
        paragraph_texts.append(paragraph_text)

    clean_text = '\n'.join(paragraph_texts).strip()

    # Extract images
    for node in cell._element.iter():
        if node.tag.endswith('}blip'):
            rId = node.get(qn('r:embed'))
            if rId:
                img_part = cell.part.related_parts[rId]
                pictures.append(img_part.blob)

    return clean_text, pictures


def read_docx_with_images(path, img_column='Photo'):
    doc = Document(path)
    rows, img_registry = [], {}
    img_col_index = None  # first column that actually contains a picture

    for t in doc.tables:
        for r_idx, row in enumerate(t.rows):
            current = []
            for c_idx, cell in enumerate(row.cells):
                txt, pics = _extract_cell(cell)

                if pics:  # keep only the *first* image per cell
                    # give every image a unique key we can round‑trip later
                    key = f"img_{len(img_registry)}"
                    img_registry[key] = pics[0]

                    # remember which column holds pictures so we can rename it
                    img_col_index = c_idx if img_col_index is None else img_col_index
                    current.append(key)
                else:
                    current.append(txt)
            rows.append(current)

    df = pd.DataFrame(rows, dtype=str)

    # treat first row as header
    df.columns = df.iloc[0]
    df = df.iloc[1:].reset_index(drop=True)

    # make sure the picture column has a nice, stable name
    if img_col_index is not None:
        df.rename(columns={df.columns[img_col_index]: img_column}, inplace=True)

    return df, img_registry


def rename_files(col_old, col_new):
    for old, new in zip(col_old, col_new):
        dir, file = os.path.split(old)
        new_path = os.path.join(dir, new)
        os.rename(old, new_path)


def create_folders(df):
    files_old = df.drop_duplicates(subset="Kennzeichen", keep="first")
    files_old = files_old["Kennzeichen"].values.tolist()
    files_old = [x for x in files_old if not x in ["", "?"]]
    for path in files_old:
        path = os.path.join(".", path)
        Path(path).mkdir(parents=True, exist_ok=True)
    return files_old


def move_files(df):
    seen = []
    for index, row in df.iterrows():
        file_ = row["Datei neu"]
        if file_ in ["", "?"]:
            continue
        file_old = row["Datei"]
        dir, _ = os.path.split(file_old)
        file_path = os.path.join(dir, file_)
        new_path = os.path.join(".", row["Kennzeichen"], file_)
        if file_ not in seen:
            shutil.move(file_path, new_path)
            seen.append(file_)


def simple_normalize(s):
    return (unicodedata.normalize("NFKD", s)
            .encode("ASCII", "ignore")
            .decode()
            .lower())


def filter_stopps(files):
    return [x for x in files for s in ["nord", "süd", "sued", "sud", "südwest", "suedwest", "sudwest", "mitte"] if
            s in simple_normalize(x)]


def move_and_rename(tuples):
    for plate, stop, file in tuples:
        path = os.path.join(".", plate)
        new_path = os.path.join(".", plate + "_" + stop)
        # move word doc to new folder
        shutil.move(file, path)
        # rename
        if os.path.exists(path):
            if not os.path.exists(new_path):
                os.rename(path, new_path)
            else:
                print("Target folder name already exists.")
        else:
            print("Original folder does not exist.")


def get_all_files_from_folder(glob_path):
    return glob.glob(glob_path)


def save_distance_sheets(paths, dfs, img_banks, img_column="Photo", max_img_size=(100, 100)):
    for file_path, df, img_registry in zip(paths, dfs, img_banks):
        _, fname = os.path.split(file_path)
        base, _ = os.path.splitext(fname)
        out_xlsx = f"{base}_Entfernung.xlsx"

        # ── make a copy so we don't mutate caller's df ───────────────────────
        df_xls = df.copy()

        # If you blank image placeholders for export, keep original keys for image insertion:
        img_keys = None
        if img_column in df_xls.columns:
            # keep the original column for mapping, but blank cells for export-data-writing if desired
            img_keys = df_xls[img_column].astype(object).where(df_xls[img_column].notna(), None).tolist()
            # Optionally blank for the exported table (so cells are empty except images)
            df_xls[img_column] = ""

        # Build list of header marker indices in the original df (rows with repeated headers)
        header_marker_idxs = df.index[df['Nr.'] == 'Nr.'].tolist()

        with pd.ExcelWriter(out_xlsx, engine='xlsxwriter') as writer:
            # Write dataframe to sheet without pandas header (we will draw custom headers)
            df_xls.to_excel(writer, index=False, sheet_name='Entfernung', header=False)
            workbook = writer.book
            worksheet = writer.sheets['Entfernung']

            # Define formats
            header_format = workbook.add_format({
                'bold': True,
                'font_color': 'white',
                'bg_color': '#294879',  # Dark blue
                'align': 'center',
                'valign': 'vcenter'
            })
            cell_format = workbook.add_format({
                'bold': True,
                'align': 'center',
                'valign': 'vcenter'
            })
            treffpunkt_format = workbook.add_format({
                'bold': True,
                'font_color': 'white',
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#294879',   # light blue background
                'font_size': 12
            })

            # ---------- Place the top header row in the correct position ----------
            # If the very first original row is a header marker (index 0), the first excel row
            # currently contains that marker; the column names should be written one row below.
            # Otherwise write header at excel row 0.
            # We'll compute an offset: how many header markers are at or before index 0? normally 0 or 1.
            # Simpler: if the first original row is header marker => header at excel row 1, else header at 0.
            first_header_row = 0 if (len(header_marker_idxs) == 0 or header_marker_idxs[0] != 0) else 1
            # write column names at the calculated position
            for col_num, value in enumerate(df_xls.columns.values):
                worksheet.write(first_header_row, col_num, value, header_format)

            # ---------- Write all data cells with the cell_format ----------
            # We must write cell values to the correct excel rows: each original row i
            # is written at excel_row = i + shift + 0, where shift = number of header markers < = i?
            # Because we used header=False, pandas wrote raw rows starting at Excel row 0, so the
            # original row i was written at excel_row = i + number_of_header_markers_before_row_i
            # (the header marker rows themselves are present in the sheet as "data rows" from df_xls).
            # We'll re-write each visible cell using cell_format to apply bold/center.
            # Note: df_xls has the same number of rows as df (we didn't insert Treffpunkt rows into df_xls here).
            # But the sheet will later receive the merged Treffpunkt rows via your existing logic;
            # to keep behavior identical, we still write all data cells here where pandas wrote them.

            # Compute a helper function to count header markers before index i
            def header_shift_for_index(i):
                # count how many header_marker_idxs are strictly less than i
                return sum(1 for h in header_marker_idxs if h < i)

            # Re-write every cell using the cell_format at the row pandas used when exporting df_xls (accounting for shift)
            for orig_idx in df_xls.index:
                shift = header_shift_for_index(orig_idx)
                excel_row = orig_idx + shift  # 0-based excel row since header=False
                for col_num, col in enumerate(df_xls.columns):
                    val = df_xls.iloc[orig_idx, col_num]
                    # don't overwrite Treffpunkt merged cell writing later; safe to rewrite data cells
                    worksheet.write(excel_row, col_num, val, cell_format)

            # ---------- Insert Treffpunkt merged header for repeating header rows ----------
            # For each header marker row in the original df, find the excel row where that marker was written,
            # then insert the merged Treffpunkt row above the actual header row (i.e. at that excel row).
            for row_num in header_marker_idxs:
                # excel row where the header marker sits (pandas exported it at row = row_num + shift)
                shift = header_shift_for_index(row_num)
                marker_excel_row = row_num + shift

                # Treffpunkt value should be taken from the row below the marker in the original df:
                next_row = row_num + 1 if (row_num + 1) in df.index else None
                treffpunkt_value = ""
                if next_row is not None and 'Treffpunkt' in df.columns:
                    treffpunkt_value = str(df.loc[next_row, 'Treffpunkt']).strip()

                if treffpunkt_value:
                    # Merge at the marker_excel_row (which is where the marker row is in sheet)
                    worksheet.merge_range(
                        marker_excel_row,
                        0,
                        marker_excel_row,
                        len(df.columns) - 1,
                        treffpunkt_value,
                        treffpunkt_format
                    )

                    # After merging and writing Treffpunkt, write the column headers on the row below merged Treffpunkt.
                    # Column header row (should align with where the marker row used to contain header names)
                    header_row_excel = marker_excel_row + 1
                    for col_num, value in enumerate(df.columns.values):
                        worksheet.write(header_row_excel, col_num, value, header_format)

            # ---------- Auto-size every non-image column ----------
            for col in df_xls.columns:
                if col == img_column:
                    continue
                width = max(df_xls[col].astype(str).map(len).max(), len(col))
                idx = df_xls.columns.get_loc(col)
                worksheet.set_column(idx, idx, width)

            # ---------- Insert pictures (correct Excel row calculation) ----------
            if img_column in df_xls.columns and img_keys is not None:
                img_col_idx = df_xls.columns.get_loc(img_column)
                max_col_char = 0  # final column width

                for orig_idx, key in enumerate(img_keys):
                    if key is None:
                        continue

                    # compute excel row where this original row was written
                    shift = header_shift_for_index(orig_idx)
                    excel_row = orig_idx + shift

                    if key not in img_registry:
                        continue

                    raw = img_registry[key]
                    buf = BytesIO(raw)
                    w_px, h_px = Image.open(buf).size

                    # scale down if needed
                    max_w, max_h = max_img_size
                    scale = min(1, max_w / w_px, max_h / h_px)

                    w_px_scaled, h_px_scaled = int(w_px * scale), int(h_px * scale)
                    row_height_pts = h_px_scaled * 0.75  # px → points
                    col_width_char = w_px_scaled / 7  # px → Excel chars
                    max_col_char = max(max_col_char, col_width_char)

                    # set row height for the sheet row where the image belongs
                    worksheet.set_row(excel_row, row_height_pts)

                    # insert the picture at the computed excel_row
                    buf.seek(0)
                    worksheet.insert_image(
                        excel_row, img_col_idx, "",
                        {
                            "image_data": buf,
                            "x_scale": scale,
                            "y_scale": scale,
                            "x_offset": 0,
                            "y_offset": 0,
                            "positioning": 1,  # moves/sizes with cells
                        },
                    )

                # final width for the picture column
                worksheet.set_column(img_col_idx, img_col_idx, max_col_char)

    # function end